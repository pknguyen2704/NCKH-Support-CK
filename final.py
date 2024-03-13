import nltk
from nltk.stem import WordNetLemmatizer

import spacy
from docx import Document
from openpyxl import Workbook

nltk.download('wordnet')
nlp = spacy.load('en_core_web_sm')
def read_document(file_path, suffix_list):
    matching_words = []
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        for word in words:
            # print(word)
            word_tmp = get_base_word(word)
            # print(word_tmp)
            for suffix in suffix_list:
                if word_tmp.endswith(suffix):
                    matching_words.append(word_tmp)
    return matching_words


def get_base_word(word):
    # Process the word with spaCy
    word = word.lower()

    if word.endswith('.') or word.endswith('?') or word.endswith(',') or word.endswith(':') or word.endswith('!') or word.endswith("'"):
        word = word[:-1]
    if word.endswith("'s"):
        word = word[:-2]
    lemmatizer = WordNetLemmatizer()
    base_word = lemmatizer.lemmatize(word)
    return base_word

def read_suffix(file_path):
    suffix_list = []
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        suffix_list.extend(paragraph.text.split())
    return suffix_list


def get_word_type(word, word_type):
    if word_type == 'NOUN':
        return 'Noun'
    elif word_type == 'VERB':
        return 'Verb'
    elif word_type == 'ADJ':
        return 'Adjective'
    elif word_type == 'ADV':
        return 'Adverb'
    else:
        return 'Unknown'


def classify_words(word_list):
    classified_words = {'Noun': set(), 'Verb': set(), 'Adjective': set(), 'Adverb': set(), 'Unknown': set()}
    for word in word_list:
        doc = nlp(word)
        word_type = doc[0].pos_
        word_typed = get_word_type(word, word_type)
        classified_words[word_typed].add(word)
    return classified_words


def create_table(classified_words, suffix_list):
    table = [[''] + list(classified_words.keys())]

    for suffix in suffix_list:
        row = [suffix]
        for word_type in classified_words.keys():
            words_with_suffix = set(word for word in classified_words[word_type] if word.endswith(suffix))
            row.append(', '.join(sorted(words_with_suffix)))
        table.append(row)

    return table


def write_to_excel(table, output_excel_path):
    workbook = Workbook()
    sheet = workbook.active

    for row in table:
        sheet.append(row)

    workbook.save(output_excel_path)


def write_statistics_by_suffix_to_excel(statistics, output_statistics_excel_path):
    workbook = Workbook()
    sheet = workbook.active

    header_row = ['', 'Noun', 'Verb', 'Adjective', 'Adverb', 'Unknown']
    sheet.append(header_row)

    # Data rows
    for suffix, data in statistics.items():
        row = [suffix] + [data[word_type] for word_type in ['Noun', 'Verb', 'Adjective', 'Adverb', 'Unknown']]
        sheet.append(row)

    workbook.save(output_statistics_excel_path)


def calculate_statistics_by_suffix(classified_words, suffix_list, word_list):
    statistics = {}
    for suffix in suffix_list:
        suffix_statistics = {'Noun': 0, 'Verb': 0, 'Adjective': 0, 'Adverb': 0, 'Unknown': 0}
        for word_type in classified_words:
            count = 0
            for word in classified_words[word_type]:
                if word.endswith(suffix):
                    for word_temp in word_list:
                        if word == word_temp:
                            count += 1
            suffix_statistics[word_type]+=count
        statistics[suffix] = suffix_statistics

    return statistics

def main():
    document_path = 'mau.docx'
    suffix_list_path = 'suffix.docx'
    result_path = 'result.xlsx'
    statistics_path = 'statistics.xlsx'

    suffix_list = read_suffix(suffix_list_path)
    matching_words = read_document(document_path, suffix_list)

    # for word in matching_words:
    #     if word == 'stimulate':
    #         print(word)
    classified_words = classify_words(matching_words)
    table = create_table(classified_words, suffix_list)
    write_to_excel(table, result_path)
    print(f"Result written to {result_path}")
    statistics = calculate_statistics_by_suffix(classified_words, suffix_list, matching_words)
    write_statistics_by_suffix_to_excel(statistics, statistics_path)
    print(f"Statistics written to {statistics_path}")

if __name__ == "__main__":
    main()
